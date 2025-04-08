#!/usr/bin/env python3
"""
Gerador de Changelog Automatizado

Este script busca issues no Jira, gera descrições de changelog usando OpenAI,
e cria um documento Word estruturado com o resultado.

Requer as seguintes variáveis de ambiente ou configs no Google Colab:
- jira_url: URL da instância do Jira
- api_token: Token de API do Jira
- email: Email para autenticação no Jira
- OPENAI_API_KEY: Chave de API da OpenAI
"""

import os
import pandas as pd
import openai
from jira import JIRA
from docx import Document
from datetime import datetime
import argparse

class ChangelogGenerator:
    def __init__(self, use_colab=False, release_version="4.1"):
        """
        Inicializa o gerador de changelog.
        
        Args:
            use_colab (bool): Se True, usa o Google Colab para obter credenciais
            release_version (str): Versão da release para o changelog
        """
        self.release_version = release_version
        self.jira_url = None
        self.jira = None
        self.setup_credentials(use_colab)
        
    def setup_credentials(self, use_colab):
        """Configura as credenciais necessárias."""
        if use_colab:
            from google.colab import userdata
            self.jira_url = userdata.get('jira_url')
            api_token = userdata.get('api_token')
            email = userdata.get('email')
            os.environ["OPENAI_API_KEY"] = userdata.get('OPENAI_API_KEY')
        else:
            self.jira_url = os.environ.get('jira_url')
            api_token = os.environ.get('api_token')
            email = os.environ.get('email')
            # Assume que OPENAI_API_KEY já está no ambiente
        
        # Configurar OpenAI
        openai.api_key = os.environ.get("OPENAI_API_KEY")
        if not openai.api_key:
            raise ValueError("A chave da API OpenAI não foi definida corretamente")
            
        # Configurar Jira
        if not all([self.jira_url, api_token, email]):
            raise ValueError("Credenciais do Jira não estão completas")
        
        self.jira = JIRA(basic_auth=(email, api_token), server=self.jira_url)
    
    def extract_jira_issues(self, jql_query=None):
        """
        Extrai issues do Jira baseado em JQL.
        
        Args:
            jql_query (str): Consulta JQL personalizada
            
        Returns:
            DataFrame: DataFrame pandas com os dados das issues
        """
        print("Extraindo issues do Jira...")
        
        if jql_query is None:
            jql_query = f'type IN standardIssueTypes() AND fixVersion = v{self.release_version} AND project NOT IN (RMP, TC, TCS) ORDER BY issuetype DESC, priority DESC'
        
        issues = self.jira.search_issues(jql_query, maxResults=200)
        
        changelog_data = []
        for issue in issues:
            issue_data = {
                'key': issue.key,
                'link': f"{self.jira_url}/browse/{issue.key}",
                'tipo': issue.fields.issuetype.name,
                'modulo': issue.fields.components[0].name if issue.fields.components else "N/A",
                'sumario': issue.fields.summary,
                'descricao': issue.fields.description
            }
            changelog_data.append(issue_data)
        
        df = pd.DataFrame(changelog_data)
        df.to_csv('issues_para_changelog.csv', index=False)
        print(f"Extraídas {len(changelog_data)} issues para o arquivo 'issues_para_changelog.csv'")
        return df
    
    def generate_changelog(self, df=None):
        """
        Gera descrições de changelog usando OpenAI.
        
        Args:
            df (DataFrame): DataFrame com dados das issues
            
        Returns:
            DataFrame: DataFrame com descrições de changelog
        """
        print("Gerando descrições de changelog com OpenAI...")
        
        if df is None:
            try:
                df = pd.read_csv('issues_para_changelog.csv')
            except Exception as e:
                raise FileNotFoundError("Não foi possível ler o arquivo 'issues_para_changelog.csv'.") from e
        
        def gerar_changelog(row):
            prompt = f"""
            Crie um texto de changelog em português (PT-BR) para a seguinte correção ou melhoria:
            Tipo: {row['tipo']}
            Módulo: {row['modulo']}
            Sumário: {row['sumario']}
            Descrição: {row['descricao']}
            O texto deve ter apenas um parágrafo. 
            O texto deve ser conciso, técnico e seguir o formato: [Verbo no passado] + [descrição da mudança].
            """
            try:
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": prompt}],
                    max_tokens=150
                )
                return response.choices[0].message.content.strip()
            except Exception as e:
                print(f"Erro ao gerar changelog para a issue {row['key']}: {e}")
                return ""
        
        df['changelog_sugerido'] = df.apply(gerar_changelog, axis=1)
        df.to_csv('changelog_para_revisao.csv', index=False)
        print(f"Geradas {len(df)} entradas de changelog no arquivo 'changelog_para_revisao.csv'")
        return df
    
    def create_document(self, df=None):
        """
        Cria documento Word com o changelog estruturado.
        
        Args:
            df (DataFrame): DataFrame com descrições de changelog
            
        Returns:
            str: Caminho para o arquivo de saída
        """
        print("Criando documento de changelog...")
        
        if df is None:
            try:
                df = pd.read_csv('changelog_para_revisao.csv')
            except Exception as e:
                raise FileNotFoundError("Não foi possível ler o arquivo 'changelog_para_revisao.csv'.") from e
        
        # Agrupar por tipo de issue e módulo
        tipos = df['tipo'].unique()
        modulos = df['modulo'].unique()
        
        # Criar documento Word
        doc = Document()
        doc.add_heading(f'Changelog - Segura {self.release_version} - {datetime.now().strftime("%d/%m/%Y")}', 0)
        
        # Organizar por tipo de issue
        for tipo in tipos:
            doc.add_heading(f'{tipo}', 1)
            
            # Subdividir por módulo
            for modulo in modulos:
                itens = df[(df['tipo'] == tipo) & (df['modulo'] == modulo)]
                if len(itens) > 0:
                    doc.add_heading(f'{modulo}', 2)
                    
                    # Adicionar cada item com formatação consistente
                    for _, item in itens.iterrows():
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{item['changelog_sugerido']} ")
                        p.add_run(f"({item['key']})").italic = True
        
        # Nome do arquivo
        arquivo_saida = f'Changelog_Release_{self.release_version}_{datetime.now().strftime("%Y%m%d")}.docx'
        
        # Salvar documento
        doc.save(arquivo_saida)
        print(f"Documento de changelog criado: {arquivo_saida}")
        return arquivo_saida
    
    def run_full_process(self, jql_query=None):
        """
        Executa o processo completo de geração de changelog.
        
        Args:
            jql_query (str): Consulta JQL personalizada
            
        Returns:
            str: Caminho para o arquivo de saída
        """
        # Extrair issues do Jira
        df = self.extract_jira_issues(jql_query)
        
        # Gerar descrições de changelog
        df = self.generate_changelog(df)
        
        # Criar documento
        return self.create_document(df)


def main():
    """Função principal que processa argumentos e executa o gerador."""
    parser = argparse.ArgumentParser(description="Gerador de Changelog Automatizado")
    parser.add_argument("--colab", action="store_true", help="Usar Google Colab para credenciais")
    parser.add_argument("--version", "-v", type=str, default="4.1", 
                        help="Versão da release para o changelog")
    parser.add_argument("--jql", type=str, help="Consulta JQL personalizada")
    parser.add_argument("--extract-only", action="store_true", 
                        help="Apenas extrair issues do Jira")
    parser.add_argument("--generate-only", action="store_true", 
                        help="Apenas gerar descrições de changelog")
    parser.add_argument("--document-only", action="store_true", 
                        help="Apenas criar documento")
    
    args = parser.parse_args()
    
    # Inicializar gerador
    generator = ChangelogGenerator(use_colab=args.colab, release_version=args.version)
    
    # Executar etapas conforme solicitado
    if args.extract_only:
        generator.extract_jira_issues(args.jql)
    elif args.generate_only:
        generator.generate_changelog()
    elif args.document_only:
        generator.create_document()
    else:
        # Processo completo
        generator.run_full_process(args.jql)


if __name__ == "__main__":
    main()
